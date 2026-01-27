"""
Generate sample data for testing the Social Support AI pipeline.
Creates:
- Bank statements (PDF)
- Resume (DOCX)
- Assets/Liabilities (Excel)
- Credit report (PDF)
"""

import os
from pathlib import Path
from datetime import datetime, timedelta
import random

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Excel generation
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# Create uploads directory
UPLOAD_DIR = Path("data/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def generate_bank_statement():
    """Generate a sample bank statement PDF."""
    filename = UPLOAD_DIR / "bank_statement.pdf"
    
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Header
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#003366'),
        spaceAfter=6,
        alignment=1
    )
    
    story.append(Paragraph("UAE National Bank", title_style))
    story.append(Paragraph("Monthly Bank Statement", styles['Heading2']))
    story.append(Spacer(1, 12))
    
    # Account info
    info_data = [
        ['Account Holder:', 'Ahmed Mohammed Al-Mansouri'],
        ['Account Number:', '123456789'],
        ['Statement Period:', f'{(datetime.now() - timedelta(days=30)).strftime("%d-%b-%Y")} to {datetime.now().strftime("%d-%b-%Y")}'],
        ['Currency:', 'AED'],
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # Transactions
    story.append(Paragraph("Transaction Details", styles['Heading3']))
    story.append(Spacer(1, 6))
    
    transactions = [
        ['Date', 'Description', 'Amount (AED)', 'Balance (AED)'],
        ['01-Jan-2026', 'Opening Balance', '', '25,000.00'],
        ['05-Jan-2026', 'Salary Deposit', '+8,500.00', '33,500.00'],
        ['08-Jan-2026', 'Electricity Bill', '-250.00', '33,250.00'],
        ['10-Jan-2026', 'Water Bill', '-150.00', '33,100.00'],
        ['12-Jan-2026', 'Grocery Shopping', '-450.00', '32,650.00'],
        ['15-Jan-2026', 'ATM Withdrawal', '-1,000.00', '31,650.00'],
        ['20-Jan-2026', 'Freelance Income', '+1,200.00', '32,850.00'],
        ['22-Jan-2026', 'Restaurant', '-180.00', '32,670.00'],
        ['25-Jan-2026', 'Transfer to Savings', '-5,000.00', '27,670.00'],
        ['28-Jan-2026', 'Closing Balance', '', '27,670.00'],
    ]
    
    trans_table = Table(transactions, colWidths=[1.2*inch, 2.5*inch, 1.5*inch, 1.8*inch])
    trans_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(trans_table)
    story.append(Spacer(1, 12))
    
    # Summary
    summary_data = [
        ['Opening Balance:', 'AED 25,000.00'],
        ['Total Credits:', 'AED 9,700.00'],
        ['Total Debits:', 'AED 7,030.00'],
        ['Closing Balance:', 'AED 27,670.00'],
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 3*inch])
    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#90EE90')),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    story.append(summary_table)
    
    doc.build(story)
    print(f"✓ Generated bank statement: {filename}")


def generate_credit_report():
    """Generate a sample credit report PDF."""
    filename = UPLOAD_DIR / "credit_report.pdf"
    
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Header
    story.append(Paragraph("Credit Report", styles['Heading1']))
    story.append(Spacer(1, 6))
    
    # Personal info
    info_data = [
        ['Name:', 'Ahmed Mohammed Al-Mansouri'],
        ['Date of Birth:', '15-Mar-1985'],
        ['Emirates ID:', '784123456789-1'],
        ['Report Date:', datetime.now().strftime("%d-%b-%Y")],
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # Credit Score
    story.append(Paragraph("Credit Score Summary", styles['Heading3']))
    credit_data = [
        ['Credit Score:', '720'],
        ['Score Range:', '300 - 900'],
        ['Rating:', 'GOOD'],
        ['Last Updated:', datetime.now().strftime("%d-%b-%Y")],
    ]
    
    credit_table = Table(credit_data, colWidths=[2*inch, 4*inch])
    credit_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 2), colors.HexColor('#E8F5E9')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ]))
    story.append(credit_table)
    story.append(Spacer(1, 12))
    
    # Credit Accounts
    story.append(Paragraph("Active Credit Accounts", styles['Heading3']))
    accounts = [
        ['Account Type', 'Creditor', 'Limit', 'Balance', 'Status'],
        ['Credit Card', 'FAB Credit Card', 'AED 50,000', 'AED 12,500', 'Active'],
        ['Personal Loan', 'Emirates Islamic Bank', 'AED 150,000', 'AED 85,000', 'Active'],
        ['Car Loan', 'Abu Dhabi Islamic Bank', 'AED 200,000', 'AED 120,000', 'Active'],
    ]
    
    acc_table = Table(accounts, colWidths=[1.2*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1*inch])
    acc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976D2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    story.append(acc_table)
    
    doc.build(story)
    print(f"✓ Generated credit report: {filename}")


def generate_resume():
    """Generate a sample resume DOCX."""
    filename = UPLOAD_DIR / "resume.docx"
    
    doc = Document()
    
    # Header
    header = doc.add_heading("Ahmed Mohammed Al-Mansouri", level=1)
    header.alignment = 1  # Center
    
    contact = doc.add_paragraph("Dubai, UAE | +971-50-123-4567 | ahmed.mansouri@email.com")
    contact.alignment = 1
    
    doc.add_paragraph()
    
    # Professional Summary
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Results-driven Sales Executive with 8+ years of experience in retail management and customer "
        "relationship management. Proven track record of exceeding sales targets and building high-performing teams. "
        "Seeking a Senior Management position to leverage expertise in business development and strategic planning."
    )
    
    # Work Experience
    doc.add_heading("Work Experience", level=2)
    
    # Job 1
    job1 = doc.add_paragraph()
    job1_run = job1.add_run("Senior Sales Manager - Al-Fardan Retail Group")
    job1_run.bold = True
    doc.add_paragraph("Jan 2022 - Present | Dubai, UAE")
    doc.add_paragraph("Managed team of 15 sales professionals achieving 125% of quarterly targets", style='List Bullet')
    doc.add_paragraph("Increased customer retention rate by 35% through improved service protocols", style='List Bullet')
    doc.add_paragraph("Developed and implemented new CRM system improving efficiency by 40%", style='List Bullet')
    
    doc.add_paragraph()
    
    # Job 2
    job2 = doc.add_paragraph()
    job2_run = job2.add_run("Sales Executive - Emaar Properties")
    job2_run.bold = True
    doc.add_paragraph("Jun 2018 - Dec 2021 | Dubai, UAE")
    doc.add_paragraph("Closed real estate deals worth AED 125 million annually", style='List Bullet')
    doc.add_paragraph("Ranked top performer for 3 consecutive years", style='List Bullet')
    doc.add_paragraph("Built and maintained relationships with 200+ high-value clients", style='List Bullet')
    
    doc.add_paragraph()
    
    # Education
    doc.add_heading("Education", level=2)
    doc.add_paragraph("Bachelor of Business Administration - United Arab Emirates University", style='List Bullet')
    doc.add_paragraph("Major: Marketing, Graduated 2018, GPA: 3.7/4.0", style='List Bullet')
    
    doc.add_paragraph()
    
    # Skills
    doc.add_heading("Skills", level=2)
    skills_text = "Sales Management • Business Development • CRM Systems • Team Leadership • Customer Relations • Strategic Planning • Negotiations • Market Analysis"
    doc.add_paragraph(skills_text)
    
    # Certifications
    doc.add_heading("Certifications", level=2)
    doc.add_paragraph("Professional Sales Certification - American Association of Professional Sales", style='List Bullet')
    doc.add_paragraph("Advanced Excel for Business Analysis - Coursera", style='List Bullet')
    
    doc.save(filename)
    print(f"✓ Generated resume: {filename}")


def generate_assets_liabilities():
    """Generate a sample assets/liabilities Excel file."""
    filename = UPLOAD_DIR / "assets_liabilities.xlsx"
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Financial Assets"
    
    # Header styling
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    # Title
    ws['A1'] = "Personal Financial Statement"
    ws['A1'].font = Font(bold=True, size=14)
    ws.merge_cells('A1:B1')
    
    ws['A2'] = f"As of: {datetime.now().strftime('%d-%b-%Y')}"
    ws['A2'].font = Font(italic=True, size=10)
    
    # Assets Section
    row = 4
    ws[f'A{row}'] = "ASSETS"
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill
    ws[f'B{row}'].fill = header_fill
    
    assets = [
        ["Cash in Bank", 27670],
        ["Savings Account", 150000],
        ["Investment Portfolio", 75000],
        ["Real Estate Property", 850000],
        ["Vehicle", 120000],
        ["Jewelry & Personal Items", 50000],
    ]
    
    row = 5
    total_assets = 0
    for asset, value in assets:
        ws[f'A{row}'] = asset
        ws[f'B{row}'] = value
        ws[f'B{row}'].number_format = '#,##0.00'
        total_assets += value
        row += 1
    
    ws[f'A{row}'] = "TOTAL ASSETS"
    ws[f'A{row}'].font = Font(bold=True)
    ws[f'B{row}'] = total_assets
    ws[f'B{row}'].font = Font(bold=True)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    
    # Liabilities Section
    row += 3
    ws[f'A{row}'] = "LIABILITIES"
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill
    ws[f'B{row}'].fill = header_fill
    
    liabilities = [
        ["Personal Loan", 85000],
        ["Car Loan", 120000],
        ["Credit Card Outstanding", 12500],
        ["Home Renovation Loan", 35000],
    ]
    
    row += 1
    total_liabilities = 0
    for liability, value in liabilities:
        ws[f'A{row}'] = liability
        ws[f'B{row}'] = value
        ws[f'B{row}'].number_format = '#,##0.00'
        total_liabilities += value
        row += 1
    
    ws[f'A{row}'] = "TOTAL LIABILITIES"
    ws[f'A{row}'].font = Font(bold=True)
    ws[f'B{row}'] = total_liabilities
    ws[f'B{row}'].font = Font(bold=True)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")
    
    # Net Worth
    row += 2
    net_worth = total_assets - total_liabilities
    ws[f'A{row}'] = "NET WORTH"
    ws[f'A{row}'].font = Font(bold=True, size=12)
    ws[f'B{row}'] = net_worth
    ws[f'B{row}'].font = Font(bold=True, size=12)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
    
    # Calculate Asset-Liability Ratio
    row += 2
    ratio = total_assets / total_liabilities if total_liabilities > 0 else 0
    ws[f'A{row}'] = "Asset-Liability Ratio"
    ws[f'B{row}'] = ratio
    ws[f'B{row}'].number_format = '0.00'
    
    # Set column widths
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 15
    
    wb.save(filename)
    print(f"✓ Generated assets/liabilities: {filename}")


def generate_applicant_data():
    """Generate a text file with applicant data for reference."""
    filename = UPLOAD_DIR / "applicant_data.txt"
    
    content = """
APPLICANT INFORMATION
=====================

Personal Details:
- Full Name: Ahmed Mohammed Al-Mansouri
- Date of Birth: 15-Mar-1985
- Emirates ID: 784123456789-1
- Nationality: UAE National
- Family Size: 4 (Wife: Fatima, Children: 2)
- Contact: +971-50-123-4567

Employment:
- Current Position: Senior Sales Manager
- Company: Al-Fardan Retail Group
- Employment Status: Employed
- Years of Experience: 8+ years
- Monthly Salary: AED 8,500
- Monthly Freelance Income: AED 1,200
- Total Monthly Income: AED 9,700

Financial Summary:
- Bank Balance: AED 27,670
- Total Assets: AED 1,272,670
- Total Liabilities: AED 252,500
- Net Worth: AED 1,020,170
- Asset-Liability Ratio: 5.04

Credit Profile:
- Credit Score: 720
- Active Loans: 3 (Personal, Car, Home Renovation)
- Credit Cards: 1 (Outstanding balance: AED 12,500)
- Payment History: Good (No defaults or missed payments)

Housing:
- Accommodation Type: Owned Property
- Location: Dubai
- Property Value: AED 850,000
- Mortgage Status: No mortgage

Documents Uploaded:
[OK] Bank Statement (bank_statement.pdf)
[OK] Credit Report (credit_report.pdf)
[OK] Resume (resume.docx)
[OK] Assets/Liabilities Statement (assets_liabilities.xlsx)

Eligibility Assessment:
- Income: AED 9,700/month (Below AED 15,000 threshold) [OK]
- Family Size: 4 (Bonus consideration) [OK]
- Employment: Employed with stable income [OK]
- Assets: Good financial position with net worth of AED 1M+ [OK]
- Credit: Good credit score of 720 [OK]

Likely Recommendation:
- Eligible for social support consideration
- Recommended for economic enablement programs
- Strong candidate for career advancement opportunities
"""
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[OK] Generated applicant reference data: {filename}")


def main():
    """Generate all sample data."""
    print("=" * 60)
    print("Generating Sample Data for Social Support AI Pipeline")
    print("=" * 60)
    print()
    
    try:
        generate_bank_statement()
        generate_credit_report()
        generate_resume()
        generate_assets_liabilities()
        generate_applicant_data()
        
        print()
        print("=" * 60)
        print("[OK] All sample data generated successfully!")
        print(f"[OK] Files saved to: {UPLOAD_DIR.absolute()}")
        print("=" * 60)
        print()
        print("You can now upload these files through the Streamlit UI to test the pipeline:")
        print("  1. bank_statement.pdf")
        print("  2. credit_report.pdf")
        print("  3. resume.docx")
        print("  4. assets_liabilities.xlsx")
        
    except Exception as e:
        print(f"[ERROR] Error generating sample data: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
