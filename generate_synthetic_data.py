"""
Generate synthetic data for multiple applicants.
Creates complete profiles with all required documents organized by applicant.
"""

import os
from pathlib import Path
from datetime import datetime, timedelta
import random
import json

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from PIL import Image as PILImage, ImageDraw, ImageFont

# Excel generation
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# Create synthetic data directory
SYNTHETIC_DIR = Path("data/synthetic")
SYNTHETIC_DIR.mkdir(parents=True, exist_ok=True)

# Sample data
APPLICANTS = [
    {
        "name": "Ahmed Mohammed Al-Mansouri",
        "email": "ahmed.mansouri@email.com",
        "phone": "+971-50-123-4567",
        "dob": "15-Mar-1985",
        "emirates_id": "784123456789-1",
        "position": "Senior Sales Manager",
        "company": "Al-Fardan Retail Group",
        "monthly_salary": 8500,
        "freelance_income": 1200,
        "family_size": 4,
        "employment_status": "Employed",
        "bank_balance": 27670,
        "assets": {
            "cash": 27670,
            "savings": 150000,
            "investments": 75000,
            "property": 850000,
            "vehicle": 120000,
            "personal": 50000,
        },
        "liabilities": {
            "personal_loan": 85000,
            "car_loan": 120000,
            "credit_card": 12500,
            "home_renovation": 35000,
        },
        "credit_score": 720,
        "years_experience": 8,
    },
    {
        "name": "Fatima Hassan Al-Ketbi",
        "email": "fatima.ketbi@email.com",
        "phone": "+971-50-234-5678",
        "dob": "22-Jul-1990",
        "emirates_id": "784234567890-2",
        "position": "Software Developer",
        "company": "Microsoft UAE",
        "monthly_salary": 0,
        "freelance_income": 0,
        "family_size": 3,
        "employment_status": "Unemployed",
        "bank_balance": 45000,
        "assets": {
            "cash": 45000,
            "savings": 200000,
            "investments": 100000,
            "property": 0,
            "vehicle": 85000,
            "personal": 35000,
        },
        "liabilities": {
            "personal_loan": 50000,
            "car_loan": 60000,
            "credit_card": 8000,
            "home_renovation": 0,
        },
        "credit_score": 750,
        "years_experience": 6,
    },
    {
        "name": "Mohammad Abdul Rahman Al-Falasi",
        "email": "mohammad.falasi@email.com",
        "phone": "+971-50-345-6789",
        "dob": "10-Dec-1988",
        "emirates_id": "784345678901-3",
        "position": "Business Analyst",
        "company": "PwC Middle East",
        "monthly_salary": 9500,
        "freelance_income": 1500,
        "family_size": 5,
        "employment_status": "Employed",
        "bank_balance": 32000,
        "assets": {
            "cash": 32000,
            "savings": 120000,
            "investments": 60000,
            "property": 750000,
            "vehicle": 95000,
            "personal": 40000,
        },
        "liabilities": {
            "personal_loan": 75000,
            "car_loan": 85000,
            "credit_card": 10000,
            "home_renovation": 45000,
        },
        "credit_score": 680,
        "years_experience": 7,
    },
    {
        "name": "Noor Saeed Al-Maktoum",
        "email": "noor.maktoum@email.com",
        "phone": "+971-50-456-7890",
        "dob": "05-Aug-1992",
        "emirates_id": "784456789012-4",
        "position": "Marketing Manager",
        "company": "Emirates NBD",
        "monthly_salary": 10500,
        "freelance_income": 800,
        "family_size": 2,
        "employment_status": "Employed",
        "bank_balance": 38000,
        "assets": {
            "cash": 38000,
            "savings": 180000,
            "investments": 90000,
            "property": 0,
            "vehicle": 110000,
            "personal": 45000,
        },
        "liabilities": {
            "personal_loan": 60000,
            "car_loan": 100000,
            "credit_card": 15000,
            "home_renovation": 0,
        },
        "credit_score": 710,
        "years_experience": 5,
    },
    {
        "name": "Khalid Ibrahim Al-Dhaheri",
        "email": "khalid.dhaheri@email.com",
        "phone": "+971-50-567-8901",
        "dob": "18-Jan-1987",
        "emirates_id": "784567890123-5",
        "position": "Operations Manager",
        "company": "DP World",
        "monthly_salary": 42500,
        "freelance_income": 13000,
        "family_size": 6,
        "employment_status": "Employed",
        "bank_balance": 55000,
        "assets": {
            "cash": 55000,
            "savings": 250000,
            "investments": 120000,
            "property": 900000,
            "vehicle": 130000,
            "personal": 60000,
        },
        "liabilities": {
            "personal_loan": 95000,
            "car_loan": 110000,
            "credit_card": 18000,
            "home_renovation": 50000,
        },
        "credit_score": 740,
        "years_experience": 9,
    },
]


def create_emirate_id_image(applicant, applicant_dir):
    """Generate a simulated Emirates ID as an image."""
    filename = applicant_dir / "emirates_id.png"
    
    # Create image with ID card design
    width, height = 850, 540
    img = PILImage.new('RGB', (width, height), color='#1a472a')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback if not available
    try:
        title_font = ImageFont.truetype("arial.ttf", 32)
        header_font = ImageFont.truetype("arial.ttf", 24)
        text_font = ImageFont.truetype("arial.ttf", 18)
        small_font = ImageFont.truetype("arial.ttf", 14)
    except:
        title_font = header_font = text_font = small_font = ImageFont.load_default()
    
    # Header
    draw.text((50, 30), "EMIRATES ID CARD", font=header_font, fill='white')
    draw.text((50, 70), "Identity Card", font=small_font, fill='#FFD700')
    
    # Separator line
    draw.line([(50, 110), (800, 110)], fill='#FFD700', width=2)
    
    # Personal Information
    y_pos = 140
    draw.text((50, y_pos), f"Name: {applicant['name']}", font=text_font, fill='white')
    y_pos += 50
    draw.text((50, y_pos), f"ID Number: {applicant['emirates_id']}", font=text_font, fill='white')
    y_pos += 50
    draw.text((50, y_pos), f"Date of Birth: {applicant['dob']}", font=text_font, fill='white')
    y_pos += 50
    draw.text((50, y_pos), f"Nationality: United Arab Emirates", font=text_font, fill='white')
    y_pos += 50
    draw.text((50, y_pos), f"Issued: 01-Jan-2023", font=text_font, fill='white')
    y_pos += 50
    draw.text((50, y_pos), f"Expiry: 01-Jan-2033", font=text_font, fill='white')
    
    # Footer
    draw.line([(50, 480), (800, 480)], fill='#FFD700', width=2)
    draw.text((50, 495), "UAE Ministry of Interior", font=small_font, fill='#FFD700')
    
    img.save(filename)
    print(f"    [OK] Generated Emirates ID: emirates_id.png")


def generate_bank_statement(applicant, applicant_dir):
    """Generate a bank statement PDF."""
    filename = applicant_dir / "bank_statement.pdf"
    
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Header
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#003366'),
        spaceAfter=6,
        alignment=1
    )
    
    story.append(Paragraph("UAE National Bank", title_style))
    story.append(Paragraph("Monthly Bank Statement", styles['Heading2']))
    story.append(Spacer(1, 12))
    
    # Account info
    info_data = [
        ['Account Holder:', applicant['name']],
        ['Account Number:', '123456789'],
        ['Statement Period:', f'{(datetime.now() - timedelta(days=30)).strftime("%d-%b-%Y")} to {datetime.now().strftime("%d-%b-%Y")}'],
        ['Currency:', 'AED'],
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # Transactions
    story.append(Paragraph("Transaction Details", styles['Heading3']))
    story.append(Spacer(1, 6))
    
    opening_balance = applicant['bank_balance']
    transactions = [
        ['Date', 'Description', 'Amount (AED)', 'Balance (AED)'],
        ['01-' + datetime.now().strftime("%b-%Y"), 'Opening Balance', '', f'{opening_balance:,.2f}'],
        ['05-' + datetime.now().strftime("%b-%Y"), 'Salary Deposit', f'+{applicant["monthly_salary"]:,.2f}', f'{opening_balance + applicant["monthly_salary"]:,.2f}'],
        ['08-' + datetime.now().strftime("%b-%Y"), 'Utilities', '-500.00', f'{opening_balance + applicant["monthly_salary"] - 500:,.2f}'],
        ['15-' + datetime.now().strftime("%b-%Y"), 'Freelance Income', f'+{applicant["freelance_income"]:,.2f}', f'{opening_balance + applicant["monthly_salary"] + applicant["freelance_income"] - 500:,.2f}'],
        ['20-' + datetime.now().strftime("%b-%Y"), 'Loan Payment', '-5000.00', f'{opening_balance + applicant["monthly_salary"] + applicant["freelance_income"] - 5500:,.2f}'],
        ['25-' + datetime.now().strftime("%b-%Y"), 'Transfer to Savings', '-10000.00', f'{opening_balance + applicant["monthly_salary"] + applicant["freelance_income"] - 15500:,.2f}'],
        ['28-' + datetime.now().strftime("%b-%Y"), 'Closing Balance', '', f'{opening_balance + applicant["monthly_salary"] + applicant["freelance_income"] - 15500:,.2f}'],
    ]
    
    trans_table = Table(transactions, colWidths=[1.2*inch, 2.5*inch, 1.5*inch, 1.8*inch])
    trans_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(trans_table)
    
    doc.build(story)
    print(f"    [OK] Generated bank statement: bank_statement.pdf")


def generate_credit_report(applicant, applicant_dir):
    """Generate a credit report PDF."""
    filename = applicant_dir / "credit_report.pdf"
    
    doc = SimpleDocTemplate(str(filename), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Header
    story.append(Paragraph("Credit Report", styles['Heading1']))
    story.append(Spacer(1, 6))
    
    # Personal info
    info_data = [
        ['Name:', applicant['name']],
        ['Date of Birth:', applicant['dob']],
        ['Emirates ID:', applicant['emirates_id']],
        ['Report Date:', datetime.now().strftime("%d-%b-%Y")],
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # Credit Score
    story.append(Paragraph("Credit Score Summary", styles['Heading3']))
    credit_data = [
        ['Credit Score:', str(applicant['credit_score'])],
        ['Score Range:', '300 - 900'],
        ['Rating:', 'GOOD' if applicant['credit_score'] > 650 else 'FAIR'],
        ['Last Updated:', datetime.now().strftime("%d-%b-%Y")],
    ]
    
    credit_table = Table(credit_data, colWidths=[2*inch, 4*inch])
    credit_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 2), colors.HexColor('#E8F5E9')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ]))
    story.append(credit_table)
    story.append(Spacer(1, 12))
    
    # Credit Accounts
    story.append(Paragraph("Active Credit Accounts", styles['Heading3']))
    accounts = [
        ['Account Type', 'Creditor', 'Limit', 'Balance', 'Status'],
        ['Credit Card', 'FAB Credit Card', 'AED 50,000', f'AED {applicant["liabilities"]["credit_card"]:,.0f}', 'Active'],
        ['Personal Loan', 'Emirates Islamic Bank', 'AED 150,000', f'AED {applicant["liabilities"]["personal_loan"]:,.0f}', 'Active'],
        ['Car Loan', 'Abu Dhabi Islamic Bank', 'AED 200,000', f'AED {applicant["liabilities"]["car_loan"]:,.0f}', 'Active'],
    ]
    
    acc_table = Table(accounts, colWidths=[1.2*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1*inch])
    acc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976D2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    story.append(acc_table)
    
    doc.build(story)
    print(f"    [OK] Generated credit report: credit_report.pdf")


def generate_resume(applicant, applicant_dir):
    """Generate a resume DOCX."""
    filename = applicant_dir / "resume.docx"
    
    doc = Document()
    
    # Header
    header = doc.add_heading(applicant['name'], level=1)
    header.alignment = 1
    
    contact = doc.add_paragraph(f"Dubai, UAE | {applicant['phone']} | {applicant['email']}")
    contact.alignment = 1
    
    doc.add_paragraph()
    
    # Professional Summary
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        f"Results-driven {applicant['position']} with {applicant['years_experience']}+ years of experience. "
        "Proven track record of exceeding targets and building high-performing teams. "
        "Seeking to leverage expertise for greater impact and organizational growth."
    )
    
    # Work Experience
    doc.add_heading("Work Experience", level=2)
    
    job1 = doc.add_paragraph()
    job1_run = job1.add_run(f"{applicant['position']} - {applicant['company']}")
    job1_run.bold = True
    doc.add_paragraph("Jan 2022 - Present | Dubai, UAE")
    doc.add_paragraph("Exceeded performance targets consistently", style='List Bullet')
    doc.add_paragraph("Led and mentored cross-functional teams", style='List Bullet')
    doc.add_paragraph("Improved operational efficiency by 35%", style='List Bullet')
    
    doc.add_paragraph()
    
    # Education
    doc.add_heading("Education", level=2)
    doc.add_paragraph("Bachelor of Business Administration - United Arab Emirates University", style='List Bullet')
    doc.add_paragraph("Graduated 2015, GPA: 3.6/4.0", style='List Bullet')
    
    doc.add_paragraph()
    
    # Skills
    doc.add_heading("Skills", level=2)
    skills_text = "Leadership • Project Management • Strategic Planning • Team Building • Analysis • Problem Solving"
    doc.add_paragraph(skills_text)
    
    doc.save(filename)
    print(f"    [OK] Generated resume: resume.docx")


def generate_assets_liabilities(applicant, applicant_dir):
    """Generate assets/liabilities Excel file."""
    filename = applicant_dir / "assets_liabilities.xlsx"
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Financial Assets"
    
    # Header styling
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    # Title
    ws['A1'] = "Personal Financial Statement"
    ws['A1'].font = Font(bold=True, size=14)
    ws.merge_cells('A1:B1')
    
    ws['A2'] = f"As of: {datetime.now().strftime('%d-%b-%Y')}"
    ws['A2'].font = Font(italic=True, size=10)
    
    # Assets Section
    row = 4
    ws[f'A{row}'] = "ASSETS"
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill
    ws[f'B{row}'].fill = header_fill
    
    assets = [
        ["Cash in Bank", applicant['assets']['cash']],
        ["Savings Account", applicant['assets']['savings']],
        ["Investment Portfolio", applicant['assets']['investments']],
        ["Real Estate Property", applicant['assets']['property']],
        ["Vehicle", applicant['assets']['vehicle']],
        ["Jewelry & Personal Items", applicant['assets']['personal']],
    ]
    
    row = 5
    total_assets = 0
    for asset, value in assets:
        ws[f'A{row}'] = asset
        ws[f'B{row}'] = value
        ws[f'B{row}'].number_format = '#,##0.00'
        total_assets += value
        row += 1
    
    ws[f'A{row}'] = "TOTAL ASSETS"
    ws[f'A{row}'].font = Font(bold=True)
    ws[f'B{row}'] = total_assets
    ws[f'B{row}'].font = Font(bold=True)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    
    # Liabilities Section
    row += 3
    ws[f'A{row}'] = "LIABILITIES"
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = header_fill
    ws[f'B{row}'].fill = header_fill
    
    liabilities = [
        ["Personal Loan", applicant['liabilities']['personal_loan']],
        ["Car Loan", applicant['liabilities']['car_loan']],
        ["Credit Card Outstanding", applicant['liabilities']['credit_card']],
        ["Home Renovation Loan", applicant['liabilities']['home_renovation']],
    ]
    
    row += 1
    total_liabilities = 0
    for liability, value in liabilities:
        ws[f'A{row}'] = liability
        ws[f'B{row}'] = value
        ws[f'B{row}'].number_format = '#,##0.00'
        total_liabilities += value
        row += 1
    
    ws[f'A{row}'] = "TOTAL LIABILITIES"
    ws[f'A{row}'].font = Font(bold=True)
    ws[f'B{row}'] = total_liabilities
    ws[f'B{row}'].font = Font(bold=True)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")
    
    # Net Worth
    row += 2
    net_worth = total_assets - total_liabilities
    ws[f'A{row}'] = "NET WORTH"
    ws[f'A{row}'].font = Font(bold=True, size=12)
    ws[f'B{row}'] = net_worth
    ws[f'B{row}'].font = Font(bold=True, size=12)
    ws[f'B{row}'].number_format = '#,##0.00'
    ws[f'B{row}'].fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
    
    # Calculate Asset-Liability Ratio
    row += 2
    ratio = total_assets / total_liabilities if total_liabilities > 0 else 0
    ws[f'A{row}'] = "Asset-Liability Ratio"
    ws[f'B{row}'] = ratio
    ws[f'B{row}'].number_format = '0.00'
    
    # Set column widths
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 15
    
    wb.save(filename)
    print(f"    [OK] Generated assets/liabilities: assets_liabilities.xlsx")


def generate_applicant_metadata(applicant, applicant_dir):
    """Generate metadata JSON for each applicant."""
    filename = applicant_dir / "metadata.json"
    
    total_assets = sum(applicant['assets'].values())
    total_liabilities = sum(applicant['liabilities'].values())
    monthly_income = applicant['monthly_salary'] + applicant['freelance_income']
    
    metadata = {
        "name": applicant['name'],
        "email": applicant['email'],
        "phone": applicant['phone'],
        "emirates_id": applicant['emirates_id'],
        "dob": applicant['dob'],
        "employment": {
            "position": applicant['position'],
            "company": applicant['company'],
            "monthly_salary": applicant['monthly_salary'],
            "freelance_income": applicant['freelance_income'],
            "total_monthly_income": monthly_income,
            "employment_status": applicant['employment_status'],
            "years_experience": applicant['years_experience'],
        },
        "family": {
            "family_size": applicant['family_size'],
        },
        "financial": {
            "bank_balance": applicant['bank_balance'],
            "total_assets": total_assets,
            "total_liabilities": total_liabilities,
            "net_worth": total_assets - total_liabilities,
            "asset_liability_ratio": round(total_assets / total_liabilities, 2) if total_liabilities > 0 else 0,
        },
        "credit": {
            "credit_score": applicant['credit_score'],
        },
        "eligibility": {
            "income_eligible": monthly_income < 15000,
            "notes": f"Monthly income: AED {monthly_income:,.0f}",
        },
        "documents": [
            "bank_statement.pdf",
            "credit_report.pdf",
            "resume.docx",
            "assets_liabilities.xlsx",
            "emirates_id.png"
        ],
        "generated_at": datetime.now().isoformat(),
    }
    
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    print(f"    [OK] Generated metadata: metadata.json")


def generate_applicant_profile(applicant):
    """Generate all documents for a single applicant."""
    # Create applicant-specific directory
    applicant_dir = SYNTHETIC_DIR / applicant['name'].replace(' ', '_')
    applicant_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n  Generating profile for: {applicant['name']}")
    print(f"  Output directory: {applicant_dir}")
    
    generate_bank_statement(applicant, applicant_dir)
    generate_credit_report(applicant, applicant_dir)
    generate_resume(applicant, applicant_dir)
    generate_assets_liabilities(applicant, applicant_dir)
    create_emirate_id_image(applicant, applicant_dir)
    generate_applicant_metadata(applicant, applicant_dir)


def main():
    """Generate synthetic data for all applicants."""
    print("=" * 70)
    print("GENERATING SYNTHETIC APPLICANT DATA")
    print("=" * 70)
    print(f"Output directory: {SYNTHETIC_DIR.absolute()}")
    print(f"Number of applicants: {len(APPLICANTS)}")
    
    for applicant in APPLICANTS:
        generate_applicant_profile(applicant)
    
    print("\n" + "=" * 70)
    print("[OK] All synthetic data generated successfully!")
    print("=" * 70)
    print(f"\nDirectory structure:")
    print(f"  data/synthetic/")
    for applicant in APPLICANTS:
        applicant_name = applicant['name'].replace(' ', '_')
        print(f"    {applicant_name}/")
        print(f"      - bank_statement.pdf")
        print(f"      - credit_report.pdf")
        print(f"      - resume.docx")
        print(f"      - assets_liabilities.xlsx")
        print(f"      - emirates_id.png")
        print(f"      - metadata.json")
    print("\nYou can now upload these files through the Streamlit UI to test the pipeline!")


if __name__ == "__main__":
    main()
